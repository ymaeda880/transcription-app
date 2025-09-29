# ------------------------------------------------------------
# 📝 議事録作成（整形済みテキスト → 議事録）— modern専用・リトライなし版
# - ③の整形結果（話者分離済みテキスト）を入力に、構造化された議事録を作成
# - プロンプトは lib/prompts.py の MINUTES_MAKER グループを使用
# - 料金計算は modern usage（input/output/total）に統一
# - .txt に加えて .docx（Word）入力にも対応
# - ✅ 生成した議事録を .txt / .docx で保存できるダウンロードボタンを追加
# - ✅ 生成結果は session_state から常時レンダリング（保存ボタン後も消えない）
# ------------------------------------------------------------
from __future__ import annotations

import time
from typing import Dict, Any
from io import BytesIO

import streamlit as st
import pandas as pd
from openai import OpenAI

# ==== .docx 読み取り／書き出し（python-docx） ====
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# ==== 共通ユーティリティ ====
from lib.prompts import MINUTES_MAKER, get_group, build_prompt
from lib.tokens import extract_tokens_from_response, debug_usage_snapshot  # modern専用
from lib.costs import estimate_chat_cost_usd  # def(model, input_tokens, output_tokens)
from config.config import DEFAULT_USDJPY

# ========================== 共通設定 ==========================
st.set_page_config(page_title="④ 議事録作成", page_icon="📝", layout="wide")
st.title("④ 議事録作成 — 整形テキストから正式議事録へ")

OPENAI_API_KEY = st.secrets.get("openai", {}).get("api_key") or st.secrets.get("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    st.error("OpenAI API Key が見つかりません。.streamlit/secrets.toml を確認してください。")
    st.stop()

client = OpenAI(api_key=OPENAI_API_KEY)

# ---- セッション初期化（表示が消えない用の保険）----
st.session_state.setdefault("minutes_final_output", "")

# ========================== モデル設定補助 ==========================
def supports_temperature(model_name: str) -> bool:
    """GPT-5系は temperature 変更不可（=1固定）。"""
    return not model_name.startswith("gpt-5")

# ========================== UI ==========================
left, right = st.columns([1, 1], gap="large")

with left:
    st.subheader("プロンプト")

    group = get_group(MINUTES_MAKER)  # ← 議事録作成用プリセット

    # --- セッション初期化 ---
    if "minutes_mandatory" not in st.session_state:
        st.session_state["minutes_mandatory"] = group.mandatory_default
    if "minutes_preset_label" not in st.session_state:
        st.session_state["minutes_preset_label"] = group.label_for_key(group.default_preset_key)
    if "minutes_preset_text" not in st.session_state:
        st.session_state["minutes_preset_text"] = group.body_for_label(st.session_state["minutes_preset_label"])
    if "minutes_extra_text" not in st.session_state:
        st.session_state["minutes_extra_text"] = ""

    # --- 必須パート（編集可） ---
    mandatory = st.text_area(
        "必ず入る部分（常にプロンプトの先頭に含まれます）",
        height=220,
        key="minutes_mandatory",
    )

    # --- 追記プリセット（選択 → 本文を自動反映） ---
    def _on_change_preset():
        st.session_state["minutes_preset_text"] = group.body_for_label(
            st.session_state["minutes_preset_label"]
        )

    st.selectbox(
        "追記プリセット",
        options=group.preset_labels(),
        index=group.preset_labels().index(st.session_state["minutes_preset_label"]),
        key="minutes_preset_label",
        help="選んだ内容が上の必須文の下に自動的に連結されます。",
        on_change=_on_change_preset,
    )

    # 選択中のプリセット本文（編集可）
    preset_text = st.text_area("（編集可）プリセット本文", height=120, key="minutes_preset_text")

    # 任意の追加指示
    extra = st.text_area("追加指示（任意）", height=88, key="minutes_extra_text")

    st.subheader("モデル設定")
    model = st.selectbox(
        "モデル",
        [
            "gpt-5",
            "gpt-5-mini",
            "gpt-5-nano",
            "gpt-4.1-mini",
            "gpt-4.1",
        ],
        index=1,
    )

    temp_supported = supports_temperature(model)
    temperature = st.slider(
        "温度（0=厳格 / 2=自由）",
        0.0, 2.0, value=0.7, step=0.1,
        disabled=not temp_supported,
        help="GPT-5 系列は temperature=1 固定です",
    )
    if not temp_supported:
        st.caption("ℹ️ GPT-5 系列は temperature を変更できません（=1固定）")

    max_completion_tokens = st.slider(
        "最大出力トークン（目安）",
        min_value=1000, max_value=40000, value=12000, step=500,
        help="長めの議事録生成なら 8,000〜12,000 程度を推奨（本版はリトライなし）。",
    )

    st.subheader("通貨換算（任意）")
    usd_jpy = st.number_input("USD/JPY", min_value=50.0, max_value=500.0, value=float(DEFAULT_USDJPY), step=0.5)

    run_btn = st.button("📝 議事録を生成", type="primary", use_container_width=True)

with right:
    st.subheader("整形済みテキスト（入力）")

    # .txt と .docx を受け付ける
    up = st.file_uploader(
        "③ページの整形結果（.txt または .docx）をアップロードするか、下の欄に貼り付けてください。",
        type=["txt", "docx"],
        accept_multiple_files=False,
    )

    if up is not None:
        # .docx の場合は python-docx で抽出、.txt はそのまま読み取り
        if up.name.lower().endswith(".docx"):
            if not HAS_DOCX:
                st.error("`.docx` を読み込むには python-docx が必要です。`pip install python-docx` を実行してください。")
            else:
                data = up.read()
                try:
                    doc = Document(BytesIO(data))
                    text_from_file = "\n".join([p.text for p in doc.paragraphs])
                except Exception as e:
                    st.error(f"Wordファイルの読み込みに失敗しました: {e}")
                    text_from_file = ""
                st.session_state["minutes_source_text"] = text_from_file
        else:
            # .txt
            raw = up.read()
            try:
                text_from_file = raw.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    text_from_file = raw.decode("cp932")
                except Exception:
                    text_from_file = raw.decode(errors="ignore")
            st.session_state["minutes_source_text"] = text_from_file

    # 既定値（セッションから）
    src = st.text_area(
        "③ページの整形結果を引き継ぐか、ここに貼り付けてください。",
        value=st.session_state.get("minutes_source_text", ""),
        height=460,
        placeholder="「③ 話者分離・整形（新）」の結果を流し込む想定です。",
    )

# ========================== 実行（モデル呼び出し：リトライなし） ==========================
if run_btn:
    if not src.strip():
        st.warning("整形済みテキストを入力してください。")
    else:
        # プロンプト組み立て（lib/prompts 共通関数）
        combined = build_prompt(
            st.session_state["minutes_mandatory"],
            st.session_state["minutes_preset_text"],
            st.session_state["minutes_extra_text"],
            src,
        )

        def call_once(prompt_text: str, out_tokens: int):
            chat_kwargs: Dict[str, Any] = dict(
                model=model,
                messages=[{"role": "user", "content": prompt_text}],
                max_completion_tokens=int(out_tokens),
            )
            if temp_supported and abs(temperature - 1.0) > 1e-9:
                chat_kwargs["temperature"] = float(temperature)
            return client.chat.completions.create(**chat_kwargs)

        t0 = time.perf_counter()
        with st.spinner("議事録を生成中…"):
            resp = call_once(combined, max_completion_tokens)

            text = ""
            finish_reason = None
            if resp and getattr(resp, "choices", None):
                try:
                    text = resp.choices[0].message.content or ""
                except Exception:
                    text = getattr(resp.choices[0], "text", "")
                try:
                    finish_reason = resp.choices[0].finish_reason
                except Exception:
                    finish_reason = None

        elapsed = time.perf_counter() - t0

        if text.strip():
            st.session_state["minutes_final_output"] = text
            if finish_reason == "length":
                st.info("finish_reason=length: 出力が上限で切れています。必要に応じて最大出力トークンを増やしてください。")
        else:
            st.warning("⚠️ モデルから空の応答が返されました。レスポンス全体を表示します。")
            try:
                st.json(resp.model_dump())
            except Exception:
                st.write(resp)

        # === トークン算出（modern専用） ===
        if 'resp' in locals():
            input_tok, output_tok, total_tok = extract_tokens_from_response(resp)
            usd = estimate_chat_cost_usd(model, input_tok, output_tok)
            jpy = (usd * usd_jpy) if usd is not None else None

            # ===== 概要テーブル =====
            metrics_data = {
                "処理時間": [f"{elapsed:.2f} 秒"],
                "入力トークン": [f"{input_tok:,}"],
                "出力トークン": [f"{output_tok:,}"],
                "合計トークン": [f"{total_tok:,}"],
                "概算 (USD/JPY)": [f"${usd:,.6f} / ¥{jpy:,.2f}" if usd is not None else "—"],
            }
            st.subheader("トークンと料金の概要")
            st.table(pd.DataFrame(metrics_data))

            # === デバッグ用：modern usage スナップショット ===
            with st.expander("🔍 トークン算出の内訳（modern usage スナップショット）"):
                try:
                    st.write(debug_usage_snapshot(getattr(resp, "usage", None)))
                except Exception as e:
                    st.write({"error": str(e)})

# ========================== 生成結果の表示 ＆ ダウンロード（常時レンダリング） ==========================
final_text = (st.session_state.get("minutes_final_output") or "").strip()

if final_text:
    st.markdown("### 📝 生成結果（Markdown 表示）")
    st.markdown(final_text)

    st.subheader("📥 議事録の保存")

    # --- TXT 保存 ---
    txt_bytes = final_text.encode("utf-8")
    st.download_button(
        label="💾 テキストで保存 (.txt)",
        data=txt_bytes,
        file_name="minutes_output.txt",
        mime="text/plain",
        use_container_width=True,
        key="dl_txt_minutes",
    )

    # --- DOCX 保存 ---
    if HAS_DOCX:
        try:
            docx_buffer = BytesIO()
            doc = Document()
            for line in final_text.splitlines():
                doc.add_paragraph(line)
            doc.save(docx_buffer)
            docx_buffer.seek(0)

            st.download_button(
                label="💾 Wordで保存 (.docx)",
                data=docx_buffer,
                file_name="minutes_output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_docx_minutes",
            )
        except Exception as e:
            st.error(f"Word 出力でエラーが発生しました: {e}")
    else:
        st.info("Word 保存には `python-docx` が必要です。`pip install python-docx` を実行してください。")
